"""External candidate intake — pasted resume → structured profile → apply.

The "external" intake path is the public-internet flow: a candidate who
has no Thrive employee_id, no Axis mailbox, no Axis calendar. They land
on the Thrive lookalike (4th persona tile), paste their resume, get a
Claude-driven match against open JDs, and apply.

Why this is a separate module
-----------------------------
The internal path (``app.tools.intake``) reads from the seeded
``STORE.candidates`` map. External candidates are *created* on the fly
when they paste a resume; they have a Profile shaped almost identically
to internal employees but with a synthetic employee_id (= their email)
so the rest of the funnel — scoring, application creation, orchestrator
state machine — does not need to branch on candidate kind.

The orchestrator branches on ``Application.candidate_kind`` only at two
points: (a) when proposing R1 slots it skips the candidate's calendar,
and (b) when waiting for confirmation it expects an inbound email reply
parsed by ``app.services.inbox_parser``, not a UI click.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import List, Optional
from uuid import uuid4

from ..models import Application, Candidate, FunnelStage, Profile
from ..services.scoring_service import score_profile
from ..store import STORE
from .intake import list_open_jds


# ---------------------------------------------------------------------------
# DTOs
# ---------------------------------------------------------------------------


@dataclass
class ExternalRecommendation:
    """One ranked job for the external recommendations page."""
    job_id: str
    job_title: str
    location: str
    match_percent: float
    matched_skills: List[str]
    missing_skills: List[str]
    rationale: str


# ---------------------------------------------------------------------------
# Resume parser
# ---------------------------------------------------------------------------


# Maps a free-text keyword found in a resume to the *canonical* JD skill
# name we want to emit. Canonical names match the seeded JD required-skill
# strings exactly so the deterministic scorer (exact-match overlap) can
# pick them up. The real implementation will replace this with a Claude
# call that returns the same canonical names against the JD rubric.
_SKILL_LEXICON: list[tuple[str, str]] = [
    ("corporate salary", "corporate salary acquisition"),
    ("casa", "CASA cross-sell"),
    ("kyc", "KYC compliance"),
    ("relationship manager", "relationship management"),
    ("relationship management", "relationship management"),
    ("portfolio", "portfolio sales"),
    ("stakeholder", "stakeholder management"),
    ("branch manager", "branch operations"),
    ("branch operations", "branch operations"),
    ("team", "team leadership"),
    ("audit", "audit & compliance"),
    ("p&l", "P&L ownership"),
    ("customer experience", "customer experience"),
    ("burgundy", "HNI relationship management"),
    ("hni", "HNI relationship management"),
    ("wealth", "portfolio advisory"),
    ("mutual fund", "mutual funds"),
    ("aif", "AIF & PMS"),
    ("credit", "credit underwriting"),
    ("financial statement", "financial statement analysis"),
    ("risk", "risk assessment"),
    ("rbi", "RBI compliance"),
    ("excel", "MS Excel modelling"),
    ("secured lending", "secured lending"),
    ("sme", "secured lending"),
    ("working capital", "secured lending"),
    ("product management", "product management"),
    ("user research", "user research"),
    ("sql", "analytics SQL"),
    ("agile", "agile delivery"),
    ("mobile", "mobile banking"),
    ("process improvement", "process improvement"),
    ("vendor", "vendor management"),
    ("client acquisition", "client acquisition"),
]

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")


def extract_text_from_upload(filename: str, data: bytes) -> str:
    """Extract plain text from a PDF, DOCX, or plain-text resume upload.

    Used by the ``/external/parse-resume-file`` endpoint so candidates can
    upload the file they already have on disk instead of pasting raw text.
    Falls back to UTF-8 decode for unknown extensions.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        from io import BytesIO
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx"):
        from io import BytesIO
        from docx import Document

        doc = Document(BytesIO(data))
        # Walk the document body in document order so paragraphs and tables
        # appear interleaved (resumes and salary slips routinely put the
        # actual numbers inside tables — pulling only `doc.paragraphs`
        # silently drops them and causes Claude to report "no data found").
        chunks: list[str] = []
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                chunks.append(txt)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    # Tab-separate so Claude / regex can parse "Basic\t72,000"
                    chunks.append("\t".join(cells))
        return "\n".join(chunks)
    # Plain text / .txt / unknown — best-effort decode.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore")


def parse_resume(resume_text: str) -> Profile:
    """Extract a structured ``Profile`` from a pasted resume.

    Real implementation will call ``claude-sonnet-4-6`` with a structured-
    output schema. The deterministic fallback below is what runs in unit
    tests (no API key) and is also used as a safety net if the Claude call
    fails — better to ship a degraded profile than to 500 the candidate.
    """
    # 1. Pull the email out — the most reliable anchor in any resume.
    email_match = _EMAIL_RE.search(resume_text or "")
    email = email_match.group(0) if email_match else "unknown@external.invalid"

    # 1b. Phone / mobile number extraction — try several patterns in order.
    # Real resumes format phone numbers in many ways, and docx/pdf extraction
    # can compress or split whitespace unpredictably.
    phone: Optional[str] = None
    _text = resume_text or ""

    _phone_patterns = [
        # (1) Labeled phone: "Phone: +91 98765 43210" / "Mobile: 9876543210"
        r"(?:Phone|Mobile|Cell|Tel|Contact|Ph)[\s.:#\-]*(\+?\d[\d\s\-().]{7,17}\d)",
        # (2) +91 followed by 10 digits (with optional separators)
        r"(\+91[\s\-]?\d{5}[\s\-]?\d{5})",
        r"(\+91[\s\-]?\d{10})",
        # (3) India mobile: bare 10-digit starting with 6-9
        r"(?<!\d)([6-9]\d{9})(?!\d)",
        # (4) Any international-looking +country + digits block
        r"(\+\d{1,3}[\s\-]?\d[\d\s\-().]{7,15}\d)",
    ]

    for pat in _phone_patterns:
        m = re.search(pat, _text, re.IGNORECASE)
        if m:
            candidate_phone = m.group(1)
            # Normalize — strip spaces/dashes/parens for storage
            digits_only = re.sub(r"[^\d+]", "", candidate_phone)
            # Must have at least 10 digits total
            if len(re.sub(r"[^\d]", "", digits_only)) >= 10:
                phone = digits_only
                break

    # 2. First non-empty line is the candidate name (almost always true on
    # any reasonable resume layout). Strip bullets / decorations.
    lines = [ln.strip() for ln in (resume_text or "").splitlines() if ln.strip()]
    name = lines[0] if lines else "External Candidate"
    name = re.sub(r"[\u2022•\-\*]+", "", name).strip()
    if not name:
        name = "External Candidate"

    # 3. Skill extraction by lexicon scan (case-insensitive substring).
    # Each lexicon entry maps a resume keyword to a canonical JD skill
    # name; we emit the canonical form so the scorer's exact-match
    # overlap picks it up.
    text_lower = (resume_text or "").lower()
    seen: set = set()
    deduped: List[str] = []
    for keyword, canonical in _SKILL_LEXICON:
        if keyword in text_lower and canonical not in seen:
            seen.add(canonical)
            deduped.append(canonical)

    # 4. Naive tenure heuristic — count year ranges like "(2017 - 2020)".
    year_pairs = re.findall(r"(\d{4})\s*[-–]\s*(\d{4}|present)", text_lower)
    tenure_years = 0.0
    for start, end in year_pairs:
        try:
            start_y = int(start)
            end_y = 2026 if end == "present" else int(end)
            tenure_years += max(0, end_y - start_y)
        except ValueError:
            continue
    if tenure_years == 0.0 and "present" in text_lower:
        tenure_years = 1.0  # at least one current role

    # 5. Location guess — find a city-shaped token after a comma.
    location = "Unknown"
    loc_match = re.search(r"·\s*([A-Z][a-zA-Z ]+),\s*India", resume_text or "")
    if loc_match:
        location = loc_match.group(1).strip()

    return Profile(
        employee_id=email,  # external candidates use email as their stable id
        name=name,
        email=email,
        current_role="External Applicant",
        current_location=location,
        tenure_years=tenure_years,
        kras=[],
        skills=deduped,
        education=None,
        last_rating=None,
        phone=phone,
    )


# ---------------------------------------------------------------------------
# Recommendation
# ---------------------------------------------------------------------------


def recommend_jobs(profile: Profile) -> List[ExternalRecommendation]:
    """Score the parsed profile against every open JD; return ranked list.

    Reuses the same ``score_profile`` the internal flow uses so internal
    and external candidates are scored on the same yardstick.
    """
    out: List[ExternalRecommendation] = []
    for jd in list_open_jds():
        result = score_profile(profile, jd)
        out.append(
            ExternalRecommendation(
                job_id=jd.id,
                job_title=jd.title,
                location=jd.location,
                match_percent=result.percent,
                matched_skills=result.matched_skills,
                missing_skills=result.missing_skills,
                rationale=result.rationale,
            )
        )
    out.sort(key=lambda r: r.match_percent, reverse=True)
    return out


# ---------------------------------------------------------------------------
# Apply
# ---------------------------------------------------------------------------


def apply_external(
    profile: Profile,
    job_id: str,
    *,
    search_query: Optional[str] = None,
    recommendation_source: Optional[str] = None,
) -> Application:
    """Create (or upsert) the external candidate, then create an
    Application stamped with ``candidate_kind='external'``.

    Idempotent on candidate creation: if a candidate with the same email
    already exists in the store we reuse it (so a returning external
    applicant who applies to a 2nd JD doesn't create a duplicate). When
    the *new* profile carries richer data than the stored one (e.g. an
    extracted compensation breakdown from a freshly uploaded salary
    slip), we merge it in so HR sees the most recent state.

    ``search_query`` and ``recommendation_source`` are stamped onto the
    Application so HR can tell whether the candidate clicked Apply from
    the natural-language search section or the AI-recommendation
    section, and can see the literal phrase they searched for.
    """
    from .intake import get_jd  # local import to avoid cycle

    jd = get_jd(job_id)

    # Upsert the candidate by email and merge richer profile data
    # forward — never *clear* an existing field with None.
    cand = STORE.get_candidate(profile.email)
    if cand is None:
        cand = Candidate(id=profile.email, profile=profile, kind="external")
        STORE.add_candidate(cand)
    else:
        merged = cand.profile.model_copy()
        # Compensation always wins if the new upload extracted something.
        if profile.compensation is not None:
            merged.compensation = profile.compensation
        # Skills / role / location: prefer the freshest non-empty value.
        if profile.skills:
            merged.skills = profile.skills
        if profile.current_role and profile.current_role != "External Applicant":
            merged.current_role = profile.current_role
        if profile.current_location and profile.current_location != "Unknown":
            merged.current_location = profile.current_location
        if profile.tenure_years and profile.tenure_years > 0:
            merged.tenure_years = profile.tenure_years
        if profile.name and profile.name != "External Candidate":
            merged.name = profile.name
        cand.profile = merged

    # Always score against the freshest profile we have on file (which
    # may include the upserted skills above).
    scoring_profile = cand.profile
    result = score_profile(scoring_profile, jd)
    app = Application(
        id=f"app-{uuid4().hex[:8]}",
        candidate_id=profile.email,
        job_id=jd.id,
        match_percent=result.percent,
        match_rationale=result.rationale,
        matched_skills=result.matched_skills,
        missing_skills=result.missing_skills,
        stage=FunnelStage.APPLIED,
        candidate_kind="external",
        source="external_intake",
        search_query=search_query,
        recommendation_source=recommendation_source,  # type: ignore[arg-type]
    )
    STORE.add_application(app)
    app.log("candidate", f"External candidate applied to {jd.title} ({jd.location})")
    if search_query:
        app.log(
            "candidate",
            f"Source: typed search '{search_query}'"
            if recommendation_source == "search"
            else f"Source: AI recommendation (had searched '{search_query}')",
        )
    app.log("agent", f"Scored external profile vs JD: {result.percent:.0f}% — {result.rationale}")
    if scoring_profile.compensation and scoring_profile.compensation.current_ctc:
        ctc = scoring_profile.compensation.current_ctc
        app.log(
            "agent",
            f"Candidate's current CTC on file: ₹{ctc:,.0f} (extracted from latest salary slip)",
        )
    return app
